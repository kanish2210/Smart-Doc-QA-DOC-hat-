import fitz
import os
import shutil
import logging
from fastapi import UploadFile
from app.models.schemas import PageContent
from app.core.config import settings

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = [".pdf", ".docx", ".doc"]


def validate_file(file: UploadFile) -> str:
    ext = os.path.splitext(file.filename.lower())[1]
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Only PDF and DOCX files are accepted. Got: {file.filename}"
        )
    return ext


def save_uploaded_file(file: UploadFile) -> str:
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    file_path = os.path.join(settings.UPLOAD_DIR, file.filename)
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    logger.info(f"Saved file: {file_path}")
    return file_path


def extract_text_from_pdf(file_path: str) -> list[PageContent]:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    pages = []
    with fitz.open(file_path) as pdf:
        if pdf.page_count == 0:
            raise ValueError("PDF has no pages.")
        for i in range(pdf.page_count):
            text = pdf[i].get_text("text").strip()
            if not text:
                continue
            pages.append(PageContent(
                page_number=i + 1,
                text=text
            ))

    if not pages:
        raise ValueError(
            "No text could be extracted. "
            "This PDF may be scanned. Use a PDF with selectable text."
        )

    logger.info(f"Extracted {len(pages)} pages from PDF.")
    return pages


def extract_text_from_docx(file_path: str) -> list[PageContent]:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    try:
        from docx import Document
        doc = Document(file_path)

        pages = []
        current_page_text = []
        page_number = 1
        para_count = 0
        paras_per_page = 20

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            current_page_text.append(text)
            para_count += 1

            if para_count >= paras_per_page:
                combined = "\n".join(current_page_text)
                if combined.strip():
                    pages.append(PageContent(
                        page_number=page_number,
                        text=combined
                    ))
                page_number += 1
                current_page_text = []
                para_count = 0

        # Add remaining text
        if current_page_text:
            combined = "\n".join(current_page_text)
            if combined.strip():
                pages.append(PageContent(
                    page_number=page_number,
                    text=combined
                ))

        # Also extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    table_text.append(row_text)

            if table_text:
                pages.append(PageContent(
                    page_number=page_number,
                    text="\n".join(table_text)
                ))
                page_number += 1

        if not pages:
            raise ValueError(
                "No text could be extracted from this DOCX file."
            )

        logger.info(f"Extracted {len(pages)} sections from DOCX.")
        return pages

    except ImportError:
        raise ValueError(
            "python-docx not installed. "
            "Run: pip install python-docx"
        )


def process_pdf(file: UploadFile) -> dict:
    ext = validate_file(file)
    file_path = save_uploaded_file(file)

    if ext == ".pdf":
        pages = extract_text_from_pdf(file_path)
    else:
        pages = extract_text_from_docx(file_path)

    return {
        "file_path": file_path,
        "pages": pages,
        "total_pages": len(pages),
        "pages_with_text": len(pages)
    }